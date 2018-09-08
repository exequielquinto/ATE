from docx import Document
from docx.shared import Inches

# ------- initial code -------

document = Document()

p = document.add_paragraph()
r = p.add_run()
r.add_text('Good Morning every body,This is my ')
picPath = 'image0.png'
r.add_picture('image0.png',width=Inches(4.0), height=Inches(3))
r.add_text(' do you like it?')
document.save('demo.docx')

# ------- improved code -------

document = Document()

p = document.add_paragraph('Picture bullet section', 'List Bullet')
p = p.insert_paragraph_before('')
r = p.add_run()
r.add_picture('image0.png',width=Inches(4.0), height=Inches(3))
p = p.insert_paragraph_before('My picture title', 'Heading 1')

document.save('demo_better.docx')