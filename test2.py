import docx

document = docx.Document()
tbl = document.add_table(rows=0, cols=1)
row_cells = tbl.add_row().cells
paragraph = row_cells[0].paragraphs[0]
run = paragraph.add_run()
run.add_text('Good Morning every body,This is my ')
run.add_picture('image0.png', width = 2800000, height = 2100000)

tbl = document.add_table(rows=1, cols=1)
row_cells = tbl.add_row().cells
paragraph = row_cells[0].paragraphs[0]
run = paragraph.add_run()
run.add_text('Good Morning every body,This is my2 ')
run.add_picture('image.png', width = 2800000, height = 2100000)


#### not working for more than 1 page

document = docx.Document()
tbl = document.add_table(rows=2, cols=1)
row_cells = tbl.add_row().cells
paragraph = row_cells[0].paragraphs[0]
run = paragraph.add_run()
run.add_text('Good Morning every body,This is my ')
run.add_picture('image0.png', width = 2800000, height = 2100000)

tbl = document.add_table(rows=3, cols=1)
row_cells = tbl.add_row().cells
paragraph = row_cells[0].paragraphs[0]
run = paragraph.add_run()
run.add_text('Good Morning every body,This is my2 ')
run.add_picture('image.png', width = 2800000, height = 2100000)


document.save("demo.docx")
print 'done'